"""
Document ingestion pipeline.

Supports: PDF, TXT, MD, CSV, DOCX

Key design decisions:
1. Chunk size 500 chars — gives embeddings enough context to capture topic.
2. Document title is prepended to every chunk's page_content BEFORE embedding —
   so the embedding itself encodes which document a chunk came from.
3. Image-based PDFs (like infographics) get OCR fallback via PyMuPDF's built-in OCR
   or are skipped with a warning.
4. Rich metadata: filename, page number, chunk index, document title.
5. Non-PDF text files (txt, md, csv) are loaded directly as single-page documents.
6. DOCX files are loaded via python-docx if available, otherwise skipped.
"""

import os
import re
import shutil
import pickle
import logging

from langchain_community.document_loaders import PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document

from .config import (
    DOCUMENTS_FOLDER, PDF_FOLDER, CHROMA_PATH, BM25_CACHE_PATH,
    OLLAMA_BASE_URL, EMBED_MODEL,
    CHUNK_SIZE, CHUNK_OVERLAP, CHUNK_SEPARATORS,
    CHROMA_COLLECTION, SUPPORTED_EXTENSIONS,
)

logger = logging.getLogger(__name__)


def _clean_title_from_filename(filename: str) -> str:
    """Derive a human-readable title from a PDF filename.
    
    Examples:
        'Intro_to_MCP.pdf' -> 'Intro to MCP'
        'SQL_Window_Functions_1_1.pdf' -> 'SQL Window Functions'
        'linux-commands.pdf' -> 'Linux Commands'
        'mcp guide.pdf' -> 'MCP Guide'
    """
    name = os.path.splitext(filename)[0]
    # Replace underscores and hyphens with spaces
    name = name.replace("_", " ").replace("-", " ")
    # Remove version numbers like 1_1, v2, etc.
    name = re.sub(r'\b\d+(\s+\d+)*\b', '', name)
    # Collapse multiple spaces
    name = re.sub(r'\s+', ' ', name).strip()
    # Title case, but preserve obvious acronyms (all-caps words)
    words = name.split()
    titled = []
    for w in words:
        if w.isupper() and len(w) > 1:
            titled.append(w)  # Keep acronyms as-is (MCP, SQL, etc.)
        else:
            titled.append(w.capitalize())
    return " ".join(titled)


def _extract_first_heading(text: str) -> str | None:
    """Try to extract the first heading/title from page text."""
    lines = text.strip().split("\n")
    for line in lines[:5]:
        cleaned = line.strip()
        if len(cleaned) > 5 and len(cleaned) < 100:
            return cleaned
    return None


def _load_single_pdf(filepath: str, filename: str) -> list[Document]:
    """Load a single PDF, handling both text and image-based PDFs."""
    try:
        loader = PyMuPDFLoader(filepath)
        pages = loader.load()
    except Exception as e:
        logger.error(f"Failed to load {filename}: {e}")
        return []

    # Filter out empty pages
    non_empty = [p for p in pages if p.page_content and len(p.page_content.strip()) > 20]

    if not non_empty:
        # Might be an image-based PDF — try OCR via PyMuPDF
        logger.warning(f"{filename}: No text extracted. Attempting OCR with PyMuPDF...")
        try:
            import fitz
            doc = fitz.open(filepath)
            ocr_pages = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")
                if not text.strip():
                    try:
                        tp = page.get_textpage_ocr(flags=fitz.TEXT_PRESERVE_WHITESPACE)
                        text = page.get_text("text", textpage=tp)
                    except Exception:
                        text = ""

                if text.strip():
                    ocr_doc = Document(
                        page_content=text.strip(),
                        metadata={
                            "source": filepath,
                            "page": page_num,
                            "filename": filename,
                        }
                    )
                    ocr_pages.append(ocr_doc)
            doc.close()
            if ocr_pages:
                logger.info(f"{filename}: OCR extracted {len(ocr_pages)} pages")
                return ocr_pages
            else:
                logger.warning(f"{filename}: No text could be extracted even with OCR. Skipping.")
                return []
        except Exception as e:
            logger.warning(f"{filename}: OCR attempt failed: {e}. Skipping.")
            return []

    # Add filename metadata to all pages
    for page in non_empty:
        page.metadata["filename"] = filename

    return non_empty


def _load_single_text_file(filepath: str, filename: str) -> list[Document]:
    """Load a plain text file (.txt, .md, .csv) as one or more Document pages.
    
    Large files are split into ~10KB logical pages to keep memory
    manageable and give meaningful page numbers in metadata.
    """
    PAGE_SIZE = 10_000  # chars per logical page
    try:
        # Try UTF-8 first, fall back to latin-1
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                with open(filepath, "r", encoding=encoding) as f:
                    full_text = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            logger.error(f"{filename}: Could not decode with any known encoding.")
            return []

        if not full_text.strip():
            logger.warning(f"{filename}: File is empty. Skipping.")
            return []

        # Split into logical pages for large files
        pages = []
        for page_num, start in enumerate(range(0, len(full_text), PAGE_SIZE)):
            page_text = full_text[start:start + PAGE_SIZE].strip()
            if page_text:
                pages.append(Document(
                    page_content=page_text,
                    metadata={
                        "source": filepath,
                        "page": page_num,
                        "filename": filename,
                    }
                ))

        return pages

    except Exception as e:
        logger.error(f"Failed to load {filename}: {e}")
        return []


def _load_single_docx(filepath: str, filename: str) -> list[Document]:
    """Load a .docx file using python-docx. Returns one Document per page-ish section."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning(
            f"{filename}: python-docx not installed. "
            f"Run 'pip install python-docx' to enable DOCX support. Skipping."
        )
        return []

    try:
        doc = DocxDocument(filepath)
        full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())

        if not full_text.strip():
            logger.warning(f"{filename}: No text content in DOCX. Skipping.")
            return []

        # Treat entire DOCX as one logical page (most DOCX files aren't huge)
        # Split into ~10KB pages for very large docs
        PAGE_SIZE = 10_000
        pages = []
        for page_num, start in enumerate(range(0, len(full_text), PAGE_SIZE)):
            page_text = full_text[start:start + PAGE_SIZE].strip()
            if page_text:
                pages.append(Document(
                    page_content=page_text,
                    metadata={
                        "source": filepath,
                        "page": page_num,
                        "filename": filename,
                    }
                ))
        return pages

    except Exception as e:
        logger.error(f"Failed to load {filename}: {e}")
        return []


def _load_document(filepath: str, filename: str) -> list[Document]:
    """Route a file to the appropriate loader based on its extension."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _load_single_pdf(filepath, filename)
    elif ext in (".txt", ".md", ".csv"):
        return _load_single_text_file(filepath, filename)
    elif ext == ".docx":
        return _load_single_docx(filepath, filename)
    else:
        logger.warning(f"{filename}: Unsupported file type '{ext}'. Skipping.")
        return []


def load_and_chunk_documents(progress_callback=None) -> list[Document]:
    """Load all supported documents from the documents folder, split into chunks.
    
    Supported formats: PDF, TXT, MD, CSV, DOCX
    
    Args:
        progress_callback: Optional callable(message: str) for progress updates.
    
    Returns:
        List of Document chunks ready for embedding.
    """
    def log(msg):
        logger.info(msg)
        if progress_callback:
            progress_callback(msg)

    if not os.path.exists(DOCUMENTS_FOLDER):
        raise FileNotFoundError(f"Documents folder not found: {DOCUMENTS_FOLDER}")

    # Discover all supported files
    all_files = [
        f for f in os.listdir(DOCUMENTS_FOLDER)
        if os.path.splitext(f)[1].lower() in SUPPORTED_EXTENSIONS
    ]
    if not all_files:
        raise FileNotFoundError(
            f"No supported files found in {DOCUMENTS_FOLDER}. "
            f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    # Group by type for logging
    by_ext = {}
    for f in all_files:
        ext = os.path.splitext(f)[1].lower()
        by_ext.setdefault(ext, []).append(f)
    type_summary = ", ".join(f"{len(v)} {k}" for k, v in sorted(by_ext.items()))
    log(f"Found {len(all_files)} documents ({type_summary})")

    # ── Phase 1: Load all documents ──
    all_pages = []
    for filename in sorted(all_files):
        filepath = os.path.abspath(os.path.join(DOCUMENTS_FOLDER, filename))
        log(f"Loading: {filename}")
        pages = _load_document(filepath, filename)
        if pages:
            log(f"  -> {len(pages)} pages loaded")
            all_pages.extend(pages)
        else:
            log(f"  -> WARNING: No content extracted from {filename}")

    if not all_pages:
        raise ValueError("No content extracted from any document!")

    log(f"Total pages loaded: {len(all_pages)}")

    # ── Phase 2: Chunk with metadata enrichment ──
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=CHUNK_SEPARATORS,
    )

    enriched_chunks = []
    for page in all_pages:
        filename = page.metadata.get("filename", "unknown")
        doc_title = _clean_title_from_filename(filename)
        page_num = page.metadata.get("page", 0)

        # Split this page into chunks
        page_chunks = splitter.split_documents([page])

        for chunk_idx, chunk in enumerate(page_chunks):
            # ── KEY: Prepend document title to chunk content ──
            original_content = chunk.page_content.strip()
            prefix = f"[Source: {filename} | Topic: {doc_title}]\n"
            enriched_content = prefix + original_content

            # Hard cap at MAX_EMBED_CHARS to stay within embedding
            # model's token limit (mxbai-embed-large: 512 tokens).
            MAX_EMBED_CHARS = 1500
            if len(enriched_content) > MAX_EMBED_CHARS:
                allowed = MAX_EMBED_CHARS - len(prefix)
                original_content = original_content[:allowed]
                enriched_content = prefix + original_content

            chunk.page_content = enriched_content

            # Rich metadata
            chunk.metadata["filename"] = filename
            chunk.metadata["doc_title"] = doc_title
            chunk.metadata["page"] = page_num
            chunk.metadata["chunk_index"] = chunk_idx
            chunk.metadata["original_content"] = original_content

            enriched_chunks.append(chunk)

    log(f"Created {len(enriched_chunks)} chunks from {len(all_files)} documents")
    return enriched_chunks


# Backward compatibility alias
def load_and_chunk_pdfs(progress_callback=None) -> list[Document]:
    """Deprecated: Use load_and_chunk_documents() instead."""
    return load_and_chunk_documents(progress_callback=progress_callback)


def build_vectorstore(chunks: list[Document], progress_callback=None) -> Chroma:
    """Build Chroma vectorstore from chunks.
    
    Embeds chunks in small batches to avoid exceeding Ollama's
    context length limit for the embedding model.
    
    Returns:
        Chroma vectorstore instance.
    """
    def log(msg):
        logger.info(msg)
        if progress_callback:
            progress_callback(msg)

    embeddings = OllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_BASE_URL)

    # Delete old DB if it exists
    if os.path.exists(CHROMA_PATH):
        log("Removing old vector database...")
        import gc
        gc.collect()  # Release any Python references to SQLite
        import time as _time
        for attempt in range(3):
            try:
                shutil.rmtree(CHROMA_PATH)
                break
            except PermissionError:
                if attempt < 2:
                    _time.sleep(1)
                else:
                    # Last resort: rename instead of delete
                    backup = CHROMA_PATH + f"_old_{int(_time.time())}"
                    os.rename(CHROMA_PATH, backup)
                    log(f"Could not delete old DB, renamed to {backup}")

    log(f"Building vector store with {len(chunks)} chunks...")
    log(f"Embedding model: {EMBED_MODEL} (this may take a few minutes on CPU)")

    # Batch embedding with per-batch error recovery.
    # Uses batches of 10 for speed; if a batch fails (context length),
    # falls back to embedding that batch's chunks one-at-a-time.
    BATCH_SIZE = 10
    vectorstore = None
    failed_chunks = 0

    for batch_start in range(0, len(chunks), BATCH_SIZE):
        batch = chunks[batch_start:batch_start + BATCH_SIZE]
        batch_num = batch_start // BATCH_SIZE + 1
        total_batches = (len(chunks) + BATCH_SIZE - 1) // BATCH_SIZE

        if batch_num % 10 == 1 or batch_num == total_batches:
            log(f"Embedding batch {batch_num}/{total_batches} (chunk {batch_start+1}-{batch_start+len(batch)}/{len(chunks)})...")

        try:
            if vectorstore is None:
                vectorstore = Chroma.from_documents(
                    documents=batch,
                    embedding=embeddings,
                    persist_directory=CHROMA_PATH,
                    collection_name=CHROMA_COLLECTION,
                )
            else:
                vectorstore.add_documents(batch)
        except Exception as e:
            # Batch failed — fall back to per-chunk embedding
            error_msg = str(e)
            logger.warning(f"Batch {batch_num} failed ({error_msg[:80]}), retrying chunks individually...")
            for j, chunk in enumerate(batch):
                chunk_idx = batch_start + j
                try:
                    if vectorstore is None:
                        vectorstore = Chroma.from_documents(
                            documents=[chunk],
                            embedding=embeddings,
                            persist_directory=CHROMA_PATH,
                            collection_name=CHROMA_COLLECTION,
                        )
                    else:
                        vectorstore.add_documents([chunk])
                except Exception:
                    # Try truncation as last resort
                    try:
                        chunk.page_content = chunk.page_content[:800]
                        vectorstore.add_documents([chunk])
                    except Exception:
                        logger.error(f"Chunk {chunk_idx} skipped (could not embed)")
                        failed_chunks += 1

    if failed_chunks:
        log(f"WARNING: {failed_chunks} chunks failed to embed")
    log(f"Vector store built ({len(chunks) - failed_chunks}/{len(chunks)} chunks embedded)")

    # ── Build and cache BM25 index ──
    log("Building BM25 index...")
    _build_bm25_cache(chunks)
    log("BM25 index cached")

    return vectorstore



def load_vectorstore() -> Chroma | None:
    """Load existing Chroma vectorstore if available."""
    if not os.path.exists(CHROMA_PATH) or not os.listdir(CHROMA_PATH):
        return None

    try:
        embeddings = OllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_BASE_URL)
        return Chroma(
            persist_directory=CHROMA_PATH,
            embedding_function=embeddings,
            collection_name=CHROMA_COLLECTION,
        )
    except Exception as e:
        logger.error(f"Failed to load vectorstore: {e}")
        return None


def _build_bm25_cache(chunks: list[Document]):
    """Build BM25 index from chunks and cache to disk.
    
    We index the ORIGINAL content (without the [Source: ...] prefix)
    so BM25 scoring is based purely on document content, not the
    artificial prefix. The prefix is only for dense embeddings.
    """
    from rank_bm25 import BM25Okapi

    # Tokenize: simple whitespace + lowercase
    corpus = []
    chunk_metadata = []
    for chunk in chunks:
        # Use original content for BM25 (without source prefix)
        text = chunk.metadata.get("original_content", chunk.page_content)
        tokens = text.lower().split()
        corpus.append(tokens)
        chunk_metadata.append({
            "page_content": chunk.page_content,
            "metadata": chunk.metadata,
        })

    bm25 = BM25Okapi(corpus)

    cache = {
        "bm25": bm25,
        "corpus": corpus,
        "chunk_metadata": chunk_metadata,
    }

    with open(BM25_CACHE_PATH, "wb") as f:
        pickle.dump(cache, f)


def load_bm25_cache():
    """Load cached BM25 index."""
    if not os.path.exists(BM25_CACHE_PATH):
        return None

    try:
        with open(BM25_CACHE_PATH, "rb") as f:
            return pickle.load(f)
    except Exception as e:
        logger.error(f"Failed to load BM25 cache: {e}")
        return None
